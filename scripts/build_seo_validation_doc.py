"""Generate a 3-column Word doc for the SEO exec to validate SEOMATE findings.

Output columns:
1. Variable (ID + name)
2. SEOMATE finding — concise verdict + key data
3. (empty, for SEO exec to fill in after manual verification in Ahrefs / Semrush / etc.)

Only includes variables where we have actual measured data — i.e.,
captures with status PASSED, FAILED, or PARTIAL. Unmeasurable captures
and unwired variables are excluded.

Run via:
    .venv/Scripts/python scripts/build_seo_validation_doc.py
Output: docs/audit-validation-pixelettetech.docx
"""
from __future__ import annotations

import json
import os
import re
from pathlib import Path

# Synchronous DB access via psycopg — no asyncio (which kept hanging
# when the script was invoked through nested shell wrappers)
import psycopg

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL

# Map status to human label
STATUS_LABEL = {
    "passed": "PASSED",
    "failed": "FAILED",
    "partial": "PARTIAL",
}


_HEADER_RE = re.compile(r"^## (P\d+-\d+)\s+—\s+(.+?)\s*$", re.MULTILINE)
_DEFN_HEADER_RE = re.compile(r"^### Step 1 — Definition\s*$", re.MULTILINE)


def extract_variable_meta(taxonomy_path: Path) -> dict[str, dict]:
    """Parse the taxonomy for variable_id -> {name, short_definition}.

    Two-pass: locate every `## P<n>-<n>` header position, then for each
    variable take the slice between its header and the next header,
    and inside that slice find the first paragraph after `### Step 1 —
    Definition`. Avoids catastrophic regex backtracking that a single
    monster pattern can trigger on a ~9k-line document.
    """
    text = taxonomy_path.read_text(encoding="utf-8")
    headers = list(_HEADER_RE.finditer(text))
    out: dict[str, dict] = {}
    for i, h in enumerate(headers):
        full_header = h.group(0)
        if "(removed" in full_header:
            continue
        var_id = h.group(1)
        name = h.group(2).strip()
        # Slice = from end of this header to start of next header (or EOF)
        slice_start = h.end()
        slice_end = headers[i + 1].start() if i + 1 < len(headers) else len(text)
        block = text[slice_start:slice_end]
        defn_match = _DEFN_HEADER_RE.search(block)
        if not defn_match:
            continue
        after_defn = block[defn_match.end():]
        # First paragraph: until blank line or next ### heading
        para = after_defn.lstrip("\n").split("\n\n", 1)[0].strip()
        # Strip stray markdown
        para = re.sub(r"\s+", " ", para)
        # Keep first sentence only
        first_sentence_match = re.search(r"^[^.!?]{20,400}[.!?]", para)
        first_sentence = (first_sentence_match.group(0) if first_sentence_match else para).strip()
        if len(first_sentence) > 240:
            first_sentence = first_sentence[:237] + "..."
        out[var_id] = {"name": name, "short_definition": first_sentence}
    return out


def _stringify_evidence_value(v):
    """Compact, readable rendering of an evidence value for prose."""
    if v is None or v == "" or v == [] or v == {}:
        return None
    if isinstance(v, float):
        return f"{round(v, 2)}"
    if isinstance(v, bool):
        return "yes" if v else "no"
    if isinstance(v, (int, str)):
        return str(v)[:80]
    if isinstance(v, list):
        if len(v) <= 6 and all(isinstance(x, (str, int, float, bool)) for x in v):
            return ", ".join(str(x) for x in v)[:80]
        return f"({len(v)} items)"
    if isinstance(v, dict):
        # Render small dicts inline
        if len(v) <= 4 and all(
            isinstance(val, (str, int, float, bool)) for val in v.values()
        ):
            return ", ".join(
                f"{k}: {_stringify_evidence_value(val)}" for k, val in v.items()
            )[:120]
        return f"({len(v)} fields)"
    return None


# Rule-text + evidence-key combinations that we know how to phrase as
# a single human sentence. Most rules already read like plain English,
# so we mostly just inline the key numbers.
def _summarise_rule(rule: dict) -> str:
    text = (rule.get("rule_text") or "").strip()
    if not text:
        return ""
    evidence = rule.get("evidence") or {}
    # Pick the single most informative evidence value to inline
    inline_keys = [
        "coverage_pct", "kw_prominence_pct", "strict_match_pct",
        "loose_match_pct", "anchor_id_pct", "matched_pattern_count",
        "well_aligned_count", "confident_pct", "matches_at_1",
        "matches_at_3", "top1_match_pct", "top3_match_pct",
        "median_density_per_1000", "median_format_count", "median_coverage_pct",
        "median_score", "median_similarity", "weighted_average",
        "reachable_pct_this_audit", "refresh_pct_last_90d", "publications_per_month_3m",
        "total_photos", "total_available", "secondary_count", "amp_page_count",
        "negative_pct", "fabrication_issues", "pages_passing", "active_months_last_12",
        "median_format_count", "classification", "rendering_pattern",
        "indexable_count", "self_canonical_count", "pages_with_internal_inbound",
        "category_count", "total_categories", "deep_chain_count", "loop_count",
        "homepage_meta_indexnow", "signals_found", "is_24x7_business",
        "non_home_total", "suffix_pct", "absent_pct",
    ]
    rendered: list[str] = []
    for k in inline_keys:
        if k in evidence:
            val = _stringify_evidence_value(evidence[k])
            if val is not None:
                rendered.append(f"{k.replace('_', ' ')}: {val}")
            if len(rendered) >= 2:
                break
    # Append the rule's pass/fail bracket
    verdict = "✓" if rule.get("passed") else "✗"
    if rendered:
        return f"{verdict} {text} — {'; '.join(rendered)}"
    return f"{verdict} {text}"


def humanise_finding(var_id: str, status: str, value: dict | None, rules: list | None) -> str:
    """Compose a clean 1-3 sentence finding for the SEO exec column.

    Strategy: each rule's rule_text already reads like plain English
    (extractors write them that way intentionally). We render each
    rule as a checkmark or cross followed by the rule text, optionally
    inlining a key number from its evidence.
    """
    label = STATUS_LABEL.get(status, status.upper())

    # No rules → fall back to a one-line "the variable is in this state" summary
    if not rules:
        return f"{label}. (No rule-level detail recorded — see audit DB.)"

    rendered = [r for r in (_summarise_rule(r) for r in rules) if r]
    if not rendered:
        return f"{label}. (Rule text not available in capture.)"

    # Cap at 4 rules to keep the cell readable; SEO exec can request
    # detail if needed.
    if len(rendered) > 4:
        rendered = rendered[:3] + [f"… ({len(rules) - 3} more rules omitted)"]

    return f"{label}.\n" + "\n".join(rendered)


def collect_latest_measured_captures() -> dict[str, dict]:
    """Most-recent capture per variable_id, filtered to PASSED/FAILED/PARTIAL only."""
    # Read DB connection params from same .env the app uses
    conn_str = os.getenv("DATABASE_URL")
    if not conn_str:
        user = os.getenv("POSTGRES_USER", "seomate")
        pw = os.getenv("POSTGRES_PASSWORD", "seomate_dev")
        host = os.getenv("POSTGRES_HOST", "localhost")
        port = os.getenv("POSTGRES_PORT", "5433")
        db = os.getenv("POSTGRES_DB", "seomate")
        conn_str = f"postgresql://{user}:{pw}@{host}:{port}/{db}"

    sql = """
        SELECT DISTINCT ON (variable_id)
            variable_id, status, value, rules
        FROM captures
        ORDER BY variable_id, captured_at DESC
    """
    out: dict[str, dict] = {}
    with psycopg.connect(conn_str) as conn:
        with conn.cursor() as cur:
            cur.execute(sql)
            for var_id, status, value, rules in cur:
                if status not in ("passed", "failed", "partial"):
                    continue
                out[var_id] = {
                    "status": status,
                    "value": value,
                    "rules": rules,
                }
    return out


def build_docx(
    meta: dict[str, dict],
    captures: dict[str, dict],
    out_path: Path,
) -> int:
    doc = Document()

    # Make the document landscape so the 3 wide columns fit comfortably
    from docx.enum.section import WD_ORIENTATION
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_w
    section.page_height = new_h
    # Smaller margins
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Title
    doc.add_heading("SEOMATE Audit — pixelettetech.com", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(
        "For SEO Exec manual verification. Compare SEOMATE's finding "
        "against Ahrefs / Semrush / GBP dashboard / manual review, "
        "then fill in column 3 with your check result. ✓ = rule passed, "
        "✗ = rule failed."
    )
    r.font.size = Pt(10)

    # Table: 3 columns
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Variable (what it measures)"
    hdr[1].text = "SEOMATE finding"
    hdr[2].text = "Manual verification (your notes)"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Sort by pillar then numeric id
    sorted_var_ids = sorted(
        captures.keys(),
        key=lambda v: (v[:2], int(v.split("-")[1])),
    )

    count = 0
    for var_id in sorted_var_ids:
        cap = captures[var_id]
        m = meta.get(var_id, {"name": "(no taxonomy entry)", "short_definition": ""})
        finding = humanise_finding(var_id, cap["status"], cap["value"], cap.get("rules"))

        row = table.add_row().cells
        # Column 1: variable id + name (bold), then definition (regular)
        p = row[0].paragraphs[0]
        p.add_run(f"{var_id} — {m['name']}").bold = True
        if m.get("short_definition"):
            p2 = row[0].add_paragraph()
            r2 = p2.add_run(m["short_definition"])
            r2.font.size = Pt(9)
            r2.italic = True
        # Column 2: SEOMATE finding (already newline-formatted)
        row[1].text = finding
        for para in row[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
        # Column 3: blank
        row[2].text = ""
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        count += 1

    # Column widths — landscape A4 has ~27cm usable
    widths = (Cm(7.0), Cm(11.5), Cm(8.0))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    doc.save(out_path)
    return count


def main():
    repo_root = Path(__file__).resolve().parents[2]
    taxonomy = repo_root / "docs" / "o1-taxonomy.md"
    out_path = repo_root / "docs" / "audit-validation-pixelettetech.docx"

    meta = extract_variable_meta(taxonomy)
    print(f"Loaded meta for {len(meta)} variables from taxonomy.")
    captures = collect_latest_measured_captures()
    print(f"Loaded {len(captures)} variables with measured data.")
    rows = build_docx(meta, captures, out_path)
    print(f"Wrote: {out_path} ({rows} rows)")


if __name__ == "__main__":
    main()
