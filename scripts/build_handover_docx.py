"""Convert HANDOVER.md to docs/SEOMATE-handover.docx.

Lightweight markdown -> docx converter. Handles only the subset of
markdown used in HANDOVER.md (headings, bold/italic, inline code,
fenced code blocks, bulleted + numbered lists, pipe tables, hr,
links, paragraphs). No external markdown library; we hand-parse so
the diff between the source .md and the output .docx is auditable.

Run from the auditor/ dir:
    .venv/Scripts/python scripts/build_handover_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
SRC = REPO_ROOT / "HANDOVER.md"
OUT = REPO_ROOT / "docs" / "SEOMATE-handover.docx"


# ─── Inline parsing ─────────────────────────────────────────────────────────


_INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*)"           # **bold**
    r"|(`(?P<code>[^`]+)`)"                # `code`
    r"|(\[(?P<linktxt>[^\]]+)\]\((?P<linkurl>[^)]+)\))"  # [text](url)
    r"|(\*(?P<italic>[^*]+)\*)"            # *italic*
)


def add_inline(paragraph, text: str) -> None:
    """Add a markdown-formatted text fragment to an existing paragraph."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group("bold"):
            r = paragraph.add_run(m.group("bold"))
            r.bold = True
        elif m.group("code"):
            r = paragraph.add_run(m.group("code"))
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif m.group("linktxt"):
            r = paragraph.add_run(m.group("linktxt"))
            r.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
            r.underline = True
        elif m.group("italic"):
            r = paragraph.add_run(m.group("italic"))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ─── Table cell shading helper ──────────────────────────────────────────────


def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# ─── Block-level rendering ──────────────────────────────────────────────────


def render_table(doc: Document, header_cells: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header
    for i, h in enumerate(header_cells):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        add_inline(p, h)
        for run in p.runs:
            run.bold = True
        shade_cell(c, "1F4E79")
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(header_cells):
                continue
            c = table.rows[r_idx].cells[c_idx]
            c.text = ""
            p = c.paragraphs[0]
            add_inline(p, cell_text)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Spacer
    doc.add_paragraph()


def render_code_block(doc: Document, lines: list[str]) -> None:
    """Render a fenced code block as a single shaded paragraph with monospace."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def render_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    """Parse a pipe-table starting at lines[start] (header row). Returns (header, rows, next_index)."""
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    # Separator at start+1 (---)
    i = start + 2
    rows: list[list[str]] = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def parse_list_item_continuation(lines: list[str], i: int) -> tuple[str, int]:
    """A list item may continue on indented lines; collect them."""
    item = lines[i].lstrip()
    # Strip the leading bullet or "N." marker
    item = re.sub(r"^(\-|\*|\d+\.)\s+", "", item)
    j = i + 1
    # No multi-line list continuation in HANDOVER.md, but safe-handle:
    while j < len(lines) and lines[j].startswith("  ") and not lines[j].lstrip().startswith(("-", "*")):
        item += " " + lines[j].strip()
        j += 1
    return item, j


def render_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.lstrip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
                i += 1
                continue
            else:
                render_code_block(doc, code_buf)
                in_code = False
                code_buf = []
                i += 1
                continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Horizontal rule
        if stripped == "---":
            render_horizontal_rule(doc)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Heading
        h_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = heading.add_run(text)
            if level == 1:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            else:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            i += 1
            continue

        # Pipe table — detect: line starts with `|`, next line is separator
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|?\s*$", lines[i + 1].strip()):
            header, rows, next_i = parse_table(lines, i)
            render_table(doc, header, rows)
            i = next_i
            continue

        # Bulleted list
        if re.match(r"^[\-\*]\s+", stripped):
            while i < len(lines) and re.match(r"^[\-\*]\s+", lines[i].lstrip()):
                item_text, i = parse_list_item_continuation(lines, i)
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, item_text)
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].lstrip()):
                item_text, i = parse_list_item_continuation(lines, i)
                p = doc.add_paragraph(style="List Number")
                add_inline(p, item_text)
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"source not found: {SRC}")

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md_text = SRC.read_text(encoding="utf-8")
    render_markdown(doc, md_text)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
